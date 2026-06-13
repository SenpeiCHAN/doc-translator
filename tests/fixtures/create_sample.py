"""生成测试用 DOCX 文档（含中英混合、表格、图片）。"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

title = doc.add_heading("Introduction to Machine Learning", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "Machine learning is a subset of artificial intelligence (AI) "
    "that enables systems to learn and improve from experience "
    "without being explicitly programmed. The core idea is to "
    "develop algorithms that can identify patterns in data and "
    "make predictions based on those patterns."
)

doc.add_heading("Key Concepts", level=2)

concepts = [
    ("Supervised Learning", "The model learns from labeled training data, "
     "mapping inputs to known outputs. Common algorithms include linear "
     "regression, support vector machines (SVM), and random forests."),
    ("Unsupervised Learning", "The model finds hidden patterns in "
     "unlabeled data. Clustering algorithms like K-means and hierarchical "
     "clustering are typical examples."),
    ("Deep Learning", "A specialized form of machine learning using "
     "multi-layer neural networks. Architectures like convolutional "
     "neural networks (CNN) and transformers have revolutionized "
     "computer vision and natural language processing."),
]

for name, desc in concepts:
    p = doc.add_paragraph()
    run_b = p.add_run(f"{name}: ")
    run_b.bold = True
    run_b.font.size = Pt(12)
    run_n = p.add_run(desc)
    run_n.font.size = Pt(11)

doc.add_heading("Performance Metrics", level=2)

table = doc.add_table(rows=4, cols=3)
table.style = "Table Grid"
headers = ["Metric", "Formula", "Best Use Case"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ["Accuracy", "(TP+TN)/(TP+TN+FP+FN)", "Balanced datasets"],
    ["Precision", "TP/(TP+FP)", "When false positives are costly"],
    ["F1 Score", "2*P*R/(P+R)", "Imbalanced datasets"],
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i + 1].cells[j].text = val

doc.add_paragraph(
    "In practice, gradient descent optimization is used to minimize "
    "the loss function. The learning rate hyperparameter controls "
    "the step size during optimization. Regularization techniques "
    "such as L1 (Lasso) and L2 (Ridge) help prevent overfitting."
)

doc.add_heading("Data Preprocessing", level=2)

doc.add_paragraph(
    "Before training a model, data must be preprocessed. This includes "
    "handling missing values, feature scaling (normalization or "
    "standardization), and encoding categorical variables using "
    "techniques like one-hot encoding or label encoding."
)

doc.save("/Users/sampuichan/Desktop/claude/doc-translator/tests/fixtures/sample.docx")
print("Test DOCX created: tests/fixtures/sample.docx")
