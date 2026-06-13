"""DOCX 输出写入器：保留格式、图片、表格。"""

from __future__ import annotations

from io import BytesIO

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from doc_translator.document import ContentElement, Document
from doc_translator.writer.base import BaseWriter

ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "both": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


class DocxWriter(BaseWriter):

    def write(self, document: Document, output_path: str) -> None:
        out = DocxDocument()

        for page in document.pages:
            for elem in page.elements:
                if elem.type == "paragraph":
                    self._write_paragraph(out, elem)
                elif elem.type == "table":
                    self._write_table(out, elem)
                elif elem.type == "image" and elem.image_data:
                    try:
                        out.add_picture(BytesIO(elem.image_data))
                    except Exception:
                        pass

        self._write_vocabulary(out, document)
        out.save(output_path)

    def _write_paragraph(self, out: DocxDocument, elem: ContentElement) -> None:
        para = out.add_paragraph()
        meta = elem.meta or {}

        align_str = meta.get("alignment")
        if align_str and align_str in ALIGN_MAP:
            para.alignment = ALIGN_MAP[align_str]

        runs_meta = meta.get("runs", [])
        translated = elem.translated_text
        orig_text = elem.text
        text_for_runs = translated if translated else orig_text

        if runs_meta:
            self._write_runs_distributed(para, runs_meta, text_for_runs)
        else:
            run = para.add_run(text_for_runs)
            run.font.name = "PingFang SC"

    def _write_runs_distributed(
        self, para, runs_meta: list[dict], full_text: str,
    ) -> None:
        text_runs = [r for r in runs_meta if not r.get("is_image")]
        if not text_runs:
            run = para.add_run(full_text)
            run.font.name = "PingFang SC"
            return

        # Distribute translated text proportionally across non-image runs
        total_len = sum(len(r.get("text", "")) for r in text_runs) or 1
        start = 0
        for i, r in enumerate(text_runs):
            orig_len = len(r.get("text", ""))
            is_last = (i == len(text_runs) - 1)
            if is_last:
                chunk = full_text[start:]
            else:
                prop = orig_len / total_len
                chunk_len = max(1, round(len(full_text) * prop))
                chunk = full_text[start:start + chunk_len]
                start += chunk_len

            run = para.add_run(chunk)
            self._apply_font(run, r.get("font", {}))

    def _apply_font(self, run, font_meta: dict) -> None:
        name = font_meta.get("name")
        run.font.name = name if name else "PingFang SC"

        size = font_meta.get("size")
        if size:
            run.font.size = Pt(int(size) / 2)

        bold = font_meta.get("bold")
        if bold:
            run.font.bold = True

        italic = font_meta.get("italic")
        if italic:
            run.font.italic = True

        color = font_meta.get("color")
        if color:
            try:
                from docx.shared import RGBColor
                run.font.color.rgb = RGBColor.from_string(color)
            except Exception:
                pass

    def _write_table(self, out: DocxDocument, elem: ContentElement) -> None:
        rows = elem.table_rows
        if not rows:
            return

        num_cols = max(len(r) for r in rows)
        table = out.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"

        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    table.cell(i, j).text = cell_text

    def _write_vocabulary(self, out: DocxDocument, document: Document) -> None:
        vocab = document.metadata.get("vocabulary")
        if not vocab:
            return

        out.add_page_break()
        out.add_heading("术语表 / Vocabulary", level=1)

        table = out.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "English Term"
        hdr[1].text = "中文翻译"
        hdr[2].text = "备注"

        for entry in vocab:
            row = table.add_row().cells
            row[0].text = entry.get("en", "")
            row[1].text = entry.get("zh", "")
            row[2].text = entry.get("notes", "")
